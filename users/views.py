from django.shortcuts import render, redirect
from .forms import OrgChartsItems_crud, Requests_crud, Enterprises_crud, Enterprises_people_crud,Orgchart_crud,Files_crud,Assets_crud,EnterprisesLocations_crud,EnterprisesEnterprises_crud,Todos_crud,EnterprisesStandardsItems_crud,EnterpriseStandards_crud
from .forms import UserRegisterForm, Actions_crud,ActionsItemsStandards_crud,FilesDir_crud
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import TOfAEnterprisesItemsStandards, RolesType,OrgChartsItems, Requests,Roles,TodosTemplate,CfServices,EnterpriseStandards,Tenants, TypeOfActions,Enterprises,Standards,ItemsStandards, Actions,SystemLog,Files, EnterprisesPeople,OrgCharts,Assets,EnterprisesEnterprises,EnterprisesLocations,EnterprisesStandardsItems,ActionsItemsStandards,ActionsAssets,Todos
import  pytz
import uuid
from datetime import datetime
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
import calendar
from dateutil.relativedelta import relativedelta

from docx import *
from docx.shared import Inches
from io import BytesIO
from docx import Document
from docxtpl import DocxTemplate

# from schedule.models import Event


from django.shortcuts import render
from django_filters.views import FilterView
from .filters import TodosFilter,StandardsItemFilter,RequestsFilter

from django.db.models import Max

# IMPORTING SECRETS
import json

secrets= None
with open("../.env/secrets.json") as secrets:
    secrets = json.load(secrets)


def termsandconditions(request):
    return render(request, 'users/termsandconditions.html')

def login(request):
   return render(request, 'users/login.html')


@login_required()
def read_session(request):
    current_tenantid =request.session['current_tenantid']
    return current_tenantid

@login_required()
def profile(request):


    # using SendGrid's Python Library
    # https://github.com/sendgrid/sendgrid-python
    import os
    from sendgrid import SendGridAPIClient
    from sendgrid.helpers.mail import Mail

    message = Mail(
        from_email='noreply@cply.eu',
        to_emails='sub@itineris.uno',
        subject='Sending with Twilio SendGrid is Fun',
        html_content='<strong>and easy to do anywhere, even with Python</strong>')
    try:
        sg = SendGridAPIClient(secrets.SendGridApiKey)
        response = sg.send(message)

    except Exception as e:
        a=1


# =============================================================================
#    from django.core.mail import send_mail
#    from django.conf import settings
#    subject = 'Thank you for registering to our site'
#    message = ' it  means a world to us '
#    email_from = 'noreply@cply.eu'
#    recipient_list = ['admin@itineris.uno',]
#    send_mail( subject, message, email_from, recipient_list )
#
#    # Forza l'invio di tutte le email nella coda
#    from django.core.management import call_command
#    #call_command('send_mail')
#
#    return render(request, 'users/profile.html')
#
# =============================================================================

# =============================================================================
#
#    doc = DocxTemplate("users/Templates/users/doc/doc.docx")
#    context = { 'company_name' : "World company" }
#    doc.render(context)
#    docx_title="TEST_DOCUMENtT.docx"
#
#    f = BytesIO()
#    doc.save(f)
#    length = f.tell()
#    f.seek(0)
#    response = HttpResponse(
#        f.getvalue(),
#        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#    )
#    response['Content-Disposition'] = 'attachment; filename=' + docx_title
#    response['Content-Length'] = length
#    return response
#
# =============================================================================


# =============================================================================
#
#
#    doc = DocxTemplate("users/templates/users/doc/doc.docx")
#    context = { 'company_name' : "World company" }
#    doc.render(context)
#    docx_title="TEST_DOCUMENtT.docx"
#
#    f = BytesIO()
#    doc.save(f)
#    length = f.tell()
#    f.seek(0)
#    response = HttpResponse(
#        f.getvalue(),
#        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#    )
#    response['Content-Disposition'] = 'attachment; filename=' + docx_title
#    response['Content-Length'] = length
#    return response
#
#
#
# =============================================================================
# =============================================================================
#    document = Document()
#    docx_title="TEST_DOCUMENT.docx"
#    # ---- Cover Letter ----
#    #document.add_picture((r'%s/static/images/my-header.png' % (settings.PROJECT_PATH)), width=Inches(4))
#    document.add_paragraph()
#    #document.add_paragraph("%s" % date.today().strftime('%B %d, %Y'))
#
#    document.add_paragraph('Dear Sir or Madam:')
#    document.add_paragraph('We are pleased to help you with your widgets.')
#    document.add_paragraph('Please feel free to contact me for any additional information.')
#    document.add_paragraph('I look forward to assisting you in this project.')
#
#    document.add_paragraph()
#    document.add_paragraph('Best regards,')
#    document.add_paragraph('Acme Specialist 1]')
#    document.add_page_break()
#
#    # Prepare document for download
#    # -----------------------------
#
#
#    f = BytesIO()
#
#    document.save(f)
#    length = f.tell()
#    f.seek(0)
#    response = HttpResponse(
#        f.getvalue(),
#        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#    )
#    response['Content-Disposition'] = 'attachment; filename=' + docx_title
#    response['Content-Length'] = length
#    return response
# =============================================================================




    return render(request, 'users/profile.html')









@login_required()
def certification(request):

    return render(request, 'users/certification.html')

@login_required()
def home(request):
#fare il controllo se i dati che arrivano sono nella lista e nel formato giusto. far girare solo gli id
#cifrare gli id e cambiare nome alle variabili in html



    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    Element_SystemLog=SystemLog(event="post senza contenuto in actions",user=request.user,ip=request.META.get('REMOTE_ADDR', None))
    Element_SystemLog.save()

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
         # request.session['current_actiontypeid']=str(TypeOfActions.objects.all().first().id)
    #return render(request, 'users/profile.html')
    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()
    #current_Actions=Actions.objects.filter(type_of_actions=request.session['current_actiontypeid'])
    # current_actions_type=TypeOfActions.objects.filter(id=request.session['current_actiontypeid']).first()


    if request.method == 'POST':
        if request.POST.get('current_tenantid') is not None:
            EnterprisesV=Enterprises.objects.filter(tenants=request.POST.get('current_tenantid'))
            current_tenant=Tenants.objects.filter(id=request.POST.get('current_tenantid')).first()
            request.session['current_tenantid']=str(current_tenant.id)
            current_enterprise=Enterprises.objects.filter(tenants=request.session['current_tenantid']).first()
            request.session['current_enterpriseid']=str(current_enterprise.id)

        if request.POST.get('current_enterpriseid') is not None:
            current_enterprise=Enterprises.objects.filter(id=request.POST.get('current_enterpriseid')).first()
            request.session['current_enterpriseid']=str(current_enterprise.id)

# =============================================================================
#         if request.POST.get('current_actiontypeid') is not None:
#            # current_Actions=Actions.objects.filter(type_of_actions=request.POST.get('current_actiontypeid'))
#             current_actions_type=TypeOfActions.objects.filter(id=request.POST.get('current_actiontypeid')).first()
#
#             current_Actions = Actions.objects.filter(type_of_actions=current_actions_type, enterprises=current_enterprise)
#             headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]
#
#             return render(request, 'users/actions.html', {'headbar':headbar,'MenuItem':"/",'title':"Aggiungi azienda",'current_actions_type':current_actions_type,'current_Actions':current_Actions, 'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
#
# =============================================================================
        if request.POST.get('current_standardid') is not None:
            return redirect('standarditems_list/?current_standardid='+request.POST.get('current_standardid'))

        if request.POST.get('current_item_menu') =='Todos':
            return redirect('todos_list')


        if request.POST.get('current_item_menu') =='Requirements':
            return redirect('requirements_list')

        if request.POST.get('current_item_menu') =='Files':
            return redirect('files')
        if request.POST.get('current_item_menu') =='Basetables':
            enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
            OrgChartsV=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id)
            AssetsV=Assets.objects.filter(enterprises__tenants=current_tenant.id)
            Enterprises_enterprisesV=EnterprisesEnterprises.objects.filter(enterprises__tenants=current_tenant.id)
            Enterprises_locationsV=EnterprisesLocations.objects.filter(enterprises__tenants=current_tenant.id)
            #messages.error(request, current_tenant.id)
            return render(request, 'users/base_tables.html',{'AssetsV':AssetsV,'Enterprises_enterprisesV':Enterprises_enterprisesV,'Enterprises_locationsV':Enterprises_locationsV,'MenuItem':"Basetables",'title':"Anagrafiche di base",'OrgChartsV':OrgChartsV,'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


        if request.POST.get('current_item_menu') =='Requests':
           return redirect('requests_dett')

        if request.POST.get('current_item_menu') =='People':
           enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
           orgchart_dett=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id )

           headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

           return render(request, 'users/people.html',{'headbar':headbar,'enterprises_peopleV':enterprises_peopleV,'orgchart_dett':orgchart_dett,'MenuItem':"People",'title':"Persone",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

        if request.POST.get('current_item_menu') =='Plan':
           enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
           orgchart_dett=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id )
           headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]
           current_month = datetime.now().month
           current_year = datetime.now().year

           base_data=datetime(current_year, current_month, 1)
           last_date = datetime.now() + relativedelta(months=+11)

           month_plan = [
                [calendar.month_name[current_month], current_month , current_year],
                [calendar.month_name[( base_data + relativedelta(months=1)).month], ( base_data + relativedelta(months=1)).month , ( base_data + relativedelta(months=1)).year],
                [calendar.month_name[( base_data + relativedelta(months=2)).month], ( base_data + relativedelta(months=2)).month , ( base_data + relativedelta(months=2)).year],
                [calendar.month_name[( base_data + relativedelta(months=3)).month], ( base_data + relativedelta(months=3)).month , ( base_data + relativedelta(months=3)).year],
                [calendar.month_name[( base_data + relativedelta(months=4)).month], ( base_data + relativedelta(months=4)).month, ( base_data + relativedelta(months=4)).year],
                [calendar.month_name[( base_data + relativedelta(months=5)).month], ( base_data + relativedelta(months=5)).month , ( base_data + relativedelta(months=5)).year],
                [calendar.month_name[( base_data + relativedelta(months=6)).month], ( base_data + relativedelta(months=6)).month , ( base_data + relativedelta(months=6)).year],
                [calendar.month_name[( base_data + relativedelta(months=7)).month], ( base_data + relativedelta(months=7)).month , ( base_data + relativedelta(months=7)).year],
                [calendar.month_name[( base_data + relativedelta(months=8)).month], ( base_data + relativedelta(months=8)).month , ( base_data + relativedelta(months=8)).year],
                [calendar.month_name[( base_data + relativedelta(months=9)).month], ( base_data + relativedelta(months=9)).month , ( base_data + relativedelta(months=9)).year],
                [calendar.month_name[( base_data + relativedelta(months=10)).month], ( base_data + relativedelta(months=10)).month , ( base_data + relativedelta(months=10)).year],
                [calendar.month_name[( base_data + relativedelta(months=11)).month], ( base_data + relativedelta(months=11)).month , ( base_data + relativedelta(months=11)).year],
            ]

           return render(request, 'users/plan.html',{'last_date':last_date.date(),'current_date':base_data.date(),'current_month':current_month,'current_year':current_year,'month_plan':month_plan,'headbar':headbar,'enterprises_peopleV':enterprises_peopleV,'orgchart_dett':orgchart_dett,'MenuItem':"Plan",'title':"Calendario",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})




        if request.POST.get('current_item_menu') =='Education':
           enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
           orgchart_dett=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id )
           headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]
           TodosTemplateV=TodosTemplate.objects.filter(type_of_todos__name='Training' )
           #TodosTemplateV=TodosTemplate.objects.all()
           return render(request, 'users/education.html',{'TodosTemplateV':TodosTemplateV,'headbar':headbar,'enterprises_peopleV':enterprises_peopleV,'orgchart_dett':orgchart_dett,'MenuItem':"Education",'title':"Formazione",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})



        if request.POST.get('current_item_menu') =='Documentation':

            orgchart_dett=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id )
            return render(request, 'users/documentations.html',{'orgchart_dett':orgchart_dett,'MenuItem':"Documentation",'title':"Documentazione sistema",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


        if request.POST.get('current_item_menu') =='Assets':
            AssetsV=Assets.objects.filter(enterprises__tenants=current_tenant.id)
            headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

            return render(request, 'users/assets.html',{'headbar':headbar,'AssetsV':AssetsV,'MenuItem':"Assets",'title':"Assets",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

        if request.POST.get('current_item_menu') =='SWOT':
            ActionsV = Actions.objects.filter(enterprises=current_enterprise.id )
            return render(request, 'users/swot.html',{'ActionsV':ActionsV,'MenuItem':"SWOT",'title':"SWOT",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


    else:
        Element_SystemLog=SystemLog(event="post senza contenuto in actions",user=request.user,ip=request.META.get('REMOTE_ADDR', None),date_time_utc=datetime.now(pytz.utc))
        Element_SystemLog.save()

    current_people=EnterprisesPeople.objects.filter(auth_user=request.user.id).first()
    TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo',enterprises_people=current_people)
    ReqV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='request',enterprises_people=current_people)
    RequestsV=Requests.objects.filter(tenant=current_tenant.id, in_charge=current_people)

    todos_count=Todos.objects.filter(enterprises=request.session['current_enterpriseid'], type="todo", state="OPEN", enterprises_people__isnull= True, items_standards__isnull= True, standards__isnull= True).count()
    open_activities_number= Todos.objects.filter(enterprises=request.session['current_enterpriseid'], type="requirement", state="OPEN").count()
    satisfied_requirements_number= Todos.objects.filter(enterprises=request.session['current_enterpriseid'], type="requirement", state_requirement="CLOSED").count()
    unsatisfied_requirements_number= Todos.objects.filter(enterprises=request.session['current_enterpriseid'], type="requirement", state_requirement="OPEN").count()

    headbar = [["Attività non assegnate", todos_count, "bg-warning"], ["Attività aperte", open_activities_number, "bg-info"],["Requisiti non soddisfatti", unsatisfied_requirements_number, "bg-danger"],["Requisiti soddisfatti", satisfied_requirements_number, "bg-success"]] # placeholder

    return render(request, 'users/home.html', {'RequestsV':RequestsV,'TodosV':TodosV,'ReqV':ReqV,'headbar':headbar,'title':"Dashboard",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def files(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()

    FilesV=Files.objects.filter(id=current_enterprise.id).order_by('-type')
    edit_or_open='Open'
    files_dett_current_folder=""

    if request.method == "GET":
            files_dett_current_folder=request.GET.get('idf')
            if request.GET.get('eoo')=='Edit':
                        edit_or_open='Edit'
            else:
                        edit_or_open='Open'

            if request.GET.get('a')=="cm":
                if request.GET.get('idf')!="":
                    files_dett_current_folder=Files.objects.filter(id=request.GET.get('idf')).first()
                    FilesV=Files.objects.filter(father=request.GET.get('idf')).order_by('-type')
                else:
                        files_dett_current_folder=""
                        FilesV=Files.objects.filter(enterprises=current_enterprise.id, father__isnull=True).order_by('-type')
                return render(request, 'users/files.html',{'edit_or_open':edit_or_open,'currentfolder':files_dett_current_folder,'title':"Archivio",'files':FilesV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


            #OPEN new file or folder

            if request.GET.get('a')=="nfa": #new file
                 enterprise_dett=Enterprises.objects.filter(id=current_enterprise.id).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 F_files_dett_crud=Files_crud(initial={'father': request.GET.get('idf'),'type': "doc",'name': "New file",'enterprises': enterprise_dett.id})
                 return render(request, 'users/files_dett.html', {'father':request.GET.get('idf'),'title':"Archivio",'F_files_dett_crud':F_files_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            if request.GET.get('a')=="nfo": #new cartella
                 enterprise_dett=Enterprises.objects.filter(id=current_enterprise.id).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})

                 F_files_dett_crud=FilesDir_crud(initial={'father': request.GET.get('idf'),'type': "folder",'name': "New folder",'enterprises': enterprise_dett.id})
                 return render(request, 'users/files_dett.html', {'father':request.GET.get('idf'),'title':"Archivio",'F_files_dett_crud':F_files_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})



            #Open or download or edit
            if request.GET.get('a')=="o": #entro nella cartella
                 enterprise_dett=Enterprises.objects.filter(id=current_enterprise.id).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})

                 if request.GET.get('idf')!="":
                     files_dett_current_folder=Files.objects.filter(id=request.GET.get('idf')).first()
                     FilesV=Files.objects.filter(father=request.GET.get('idf')).order_by('-type')
                 else:
                         files_dett_current_folder=""
                         FilesV=Files.objects.filter(enterprises=current_enterprise.id,father__isnull=True).order_by('-type')




                 if edit_or_open=='Edit':
                     if Files.objects.filter(id=request.GET.get('idf')).first().type=="folder":
                         files_dett=Files.objects.filter(id=request.GET.get('idf')).first()
                         F_files_dett_crud=Files_crud(initial={'date_time_create': files_dett.date_time_create,'type': files_dett.type,'enterprises': files_dett.enterprises,'name': files_dett.name,'id': files_dett.id})
                         return render(request, 'users/files_dett.html', {'title':"Archivio",'F_files_dett_crud':F_files_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
                     else:
                         files_dett=Files.objects.filter(id=request.GET.get('idf')).first()
                         F_files_dett_crud=Files_crud(initial={'date_time_create': files_dett.date_time_create,'type': files_dett.type,'enterprises': files_dett.enterprises,'name': files_dett.name,'id': files_dett.id})
                         return render(request, 'users/files_dett.html', {'title':"Archivio",'F_files_dett_crud':F_files_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
                 else:
                     if Files.objects.filter(id=request.GET.get('idf')).first().type=="folder":
                         files_dett=Files.objects.filter(father=request.GET.get('idf')).order_by('-type')
                         files_dett_current_folder=Files.objects.filter(id=request.GET.get('idf')).first()
                         return render(request, 'users/files.html',{'edit_or_open':edit_or_open,'currentfolder':files_dett_current_folder,'title':"Archivio",'files':files_dett,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
                     else:
                         file_obj = get_object_or_404(Files, pk=request.GET.get('idf'))
                         response = HttpResponse(file_obj.content, content_type='application/octet-stream')
                         response['Content-Disposition'] = f'attachment; filename="{file_obj.name}"'
                         return response







    if request.method == "POST":

                    files_dett=Files.objects.filter(id= request.POST.get('id')).first()

                    form = Files_crud(request.POST, request.FILES, instance=files_dett)


                    if form.is_valid():
                         if request.POST.get('azione')=="Save":
                             instance=form.save(commit=False)

                             if request.POST.get('type')!="folder":
                                 if request.FILES.get('content'):
                                    instance.content = request.FILES['content'].read()
                                    instance.date_time_create=datetime.now()
                                    if instance.name == 'New file':
                                        instance.name = request.FILES['content'].name

                             instance.save()
                             files_dett=Files.objects.filter(id=request.POST.get('id')).first()


                             messages.success(request, "Salvato con successo")

                             if files_dett.father is not None:

                                 files_dett_current_folder=Files.objects.filter(id=files_dett.father.id).first()

                                 FilesV=Files.objects.filter(father=files_dett_current_folder.id).order_by('-type')

                                 return redirect('/files/?idf='+str(files_dett_current_folder.id)+'&a=o&eoo='+edit_or_open)
                             else:
                                     files_dett_current_folder=""

                                     FilesV=Files.objects.filter(enterprises=current_enterprise.id).order_by('-type')

                                     return redirect('/files/?idf=&a=cm&eoo='+edit_or_open)

                             return redirect('/files/?idf='+str(files_dett_current_folder.id)+'&a=o&eoo='+edit_or_open)

                         if request.POST.get('azione')=="Delete":
                           files_dett.delete()
                           messages.success(request, "Cancellato con successo")


                    for field, error_list in form.errors.items():
                        for error in error_list:
                            messages.error(request, error)

    FilesV=Files.objects.filter(enterprises=current_enterprise.id, father__isnull=True).order_by('-type')

    return render(request, 'users/files.html',{'edit_or_open':edit_or_open,'currentfolder': files_dett_current_folder,'title':"Archivio",'files':FilesV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def download_file(request):
    file_obj = get_object_or_404(Files, pk=request.GET.get('idf'))

    response = HttpResponse(file_obj.content, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{file_obj.name}"'

    return response

@login_required()
def base_tables(request):

    return render(request, 'users/vuota.html')

@login_required()
def about(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    return render(request, 'users/about.html', {'title':"About", 'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def actions_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('idac') is not None:
            actions_dett=Actions.objects.filter(id=request.GET.get('idac')).first()
            enterprise_dett=Enterprises.objects.filter(id=actions_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':enterprise_dett})
            #da copiare anche per ins save etc etc
            EnterprisesStandardsItemsV=EnterprisesStandardsItems.objects.filter(enterprises=enterprise_dett)
            ActionsItemsStandardsV=ActionsItemsStandards.objects.filter(actions=actions_dett)
            ActionsAssetsV=ActionsAssets.objects.filter(actions=actions_dett)
            standardV=Standards.objects.all()
            ItemsStandardsV=ItemsStandards.objects.all()
            assetsV=Assets.objects.filter(enterprises=enterprise_dett)
            F_actions_crud=Actions_crud(instance=actions_dett)
            TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo')

            return render(request, 'users/actions_dett.html', {'TodosV':TodosV,'ActionsAssetsV':ActionsAssetsV,'assetsV':assetsV,'ItemsStandardsV':ItemsStandardsV,'standardV':standardV,'ItemsStandardsV':ItemsStandardsV,'EnterprisesStandardsItemsV':EnterprisesStandardsItemsV,'ActionsItemsStandardsV':ActionsItemsStandardsV,'actions_dett':actions_dett,'title':"Azioni", 'F_actions_crud':F_actions_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=current_enterprise.id).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})

                 form = Actions_crud(request.POST)
                 if form.is_valid():
                     form.instance.id =str(uuid.uuid4())

                     #filtered_queryset = Actions.objects.filter(type_of_actions=form.instance.type_of_actions)
                     max_prot=Actions.objects.aggregate(Max('prot'))['prot__max']
                     form.instance.prot= max_prot+1
                     form.save()  # Salva i dati nel database
                     messages.success(request, "Aaggiunto con successo")
                     #messages.success(request, request.POST)


                 for field, error_list in form.errors.items():
                    for error in error_list:
                        field_label = form[field].label
                        messages.error(request, f"Errore nel campo '{field_label}': {error}")


                 actions_dett=Actions.objects.filter(id=form.instance.id).first()
                 EnterprisesStandardsItemsV=EnterprisesStandardsItems.objects.filter(enterprises=enterprise_dett)
                 ActionsItemsStandardsV=ActionsItemsStandards.objects.filter(actions=actions_dett)
                 ActionsAssetsV=ActionsAssets.objects.filter(actions=actions_dett)
                 standardV=Standards.objects.all()
                 ItemsStandardsV=ItemsStandards.objects.all()
                 assetsV=Assets.objects.filter(enterprises=enterprise_dett)
                 F_actions_crud=Actions_crud(instance=actions_dett)
                 TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo')

                 return render(request, 'users/actions_dett.html', {'TodosV':TodosV,'ActionsAssetsV':ActionsAssetsV,'assetsV':assetsV,'ItemsStandardsV':ItemsStandardsV,'standardV':standardV,'ItemsStandardsV':ItemsStandardsV,'EnterprisesStandardsItemsV':EnterprisesStandardsItemsV,'ActionsItemsStandardsV':ActionsItemsStandardsV,'actions_dett':actions_dett,'title':"Azioni", 'F_actions_crud':F_actions_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            actions_dett=Actions.objects.filter(id=request.POST.get('id')).first()

            enterprise_dett=Enterprises.objects.filter(id=actions_dett.enterprises.id).first()

            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':actions_dett})
            form = Actions_crud(request.POST, instance=actions_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Export Word":
                   actions_dett=Actions.objects.filter(id=request.POST.get('id')).first()

                   doc = DocxTemplate(actions_dett.type_of_actions.export_file)
                   #doc = DocxTemplate("users/templates/users/doc/doc.docx")


                   import html2text
                   h = html2text.HTML2Text()
                   h.ignore_links = False
                   formatted_text = h.handle(actions_dett.description)


                   context = { 'company_name' : actions_dett.enterprises , 'description' : formatted_text }
                   doc.render(context)
                   docx_title=actions_dett.name+".docx"

                   f = BytesIO()
                   doc.save(f)
                   length = f.tell()
                   f.seek(0)
                   response = HttpResponse(
                       f.getvalue(),
                       content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                   )
                   response['Content-Disposition'] = 'attachment; filename=' + docx_title
                   response['Content-Length'] = length
                   return response




                if request.POST.get('azione')=="Save":
                    form.save()
                    actions_dett=Actions.objects.filter(id=request.POST.get('id')).first()
                    F_actions_crud=Actions_crud(instance=actions_dett)
                    messages.success(request, "Salvato con successo")

                    EnterprisesStandardsItemsV=EnterprisesStandardsItems.objects.filter(enterprises=enterprise_dett)
                    ActionsItemsStandardsV=ActionsItemsStandards.objects.filter(actions=actions_dett)
                    ActionsAssetsV=ActionsAssets.objects.filter(actions=actions_dett)
                    standardV=Standards.objects.all()
                    ItemsStandardsV=ItemsStandards.objects.all()
                    assetsV=Assets.objects.filter(enterprises=enterprise_dett)
                    F_actions_crud=Actions_crud(instance=actions_dett)
                    TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo')

                    return render(request, 'users/actions_dett.html', {'TodosV':TodosV,'ActionsAssetsV':ActionsAssetsV,'assetsV':assetsV,'ItemsStandardsV':ItemsStandardsV,'standardV':standardV,'ItemsStandardsV':ItemsStandardsV,'EnterprisesStandardsItemsV':EnterprisesStandardsItemsV,'ActionsItemsStandardsV':ActionsItemsStandardsV,'actions_dett':actions_dett,'title':"Azioni", 'F_actions_crud':F_actions_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  actions_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=Actions.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Azioni",'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


    enterprises_peopleV=Actions.objects.filter(enterprises__tenants=current_tenant.id)
    TypeOfActions_dett=TypeOfActions.objects.filter(id=request.GET.get('idcat')).first()
    current_people=EnterprisesPeople.objects.filter(auth_user=request.user.id).first()
    F_actions_crud=Actions_crud(initial={'whoami': current_people.id,'enterprises': current_enterprise, 'date_actions': datetime.now(),'type_of_actions': TypeOfActions_dett.id})

    return render(request, 'users/actions_dett.html', {'cat':"mettere tipo azione",'title':"Aggiungi",'F_actions_crud':F_actions_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


@login_required()
def enterprises_dett(request):

    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()



    if request.method == "GET": #visualizzo
        if request.GET.get('iden') is not None:
            enterprise_dett=Enterprises.objects.filter(id=request.GET.get('iden')).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})
            F_enterprises_dett_crud=Enterprises_crud(instance=enterprise_dett)
            return render(request, 'users/enterprises_dett.html', {'title':"Aggiungi azienda",'F_enterprises_dett_crud':F_enterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST": #visualizzo
            if request.POST.get('id')=="":
                 a=Enterprises(id=str(uuid.uuid4()),name=request.POST.get('name'),tenants=current_tenant)
                 a.save()
                 #return render(request, 'users/vuota.html',{'a':str(uuid.uuid4())})

                 enterprise_dett=Enterprises.objects.filter(id=a.id).first()
                 F_enterprises_dett_crud=Enterprises_crud(instance=enterprise_dett)

                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/enterprises_dett.html', {'title':"Azienda", 'F_enterprises_dett_crud':F_enterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            oggetto = Enterprises.objects.get(id=request.POST.get('id'))

            if oggetto not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})


            form = Enterprises_crud(request.POST, instance=oggetto)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    enterprise_dett=Enterprises.objects.filter(id=request.POST.get('id')).first()
                    F_enterprises_dett_crud=Enterprises_crud(instance=enterprise_dett)
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/enterprises_dett.html', {'title':"Azienda",'F_enterprises_dett_crud':F_enterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  oggetto.delete()
                  messages.success(request, "Cancellato con successo")
                  return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, f"Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    F_enterprises_dett_crud=Enterprises_crud(instance=enterprise_dett)
    return render(request, 'users/enterprises_dett.html', {'title':"Aggiungi azienda",'F_enterprises_dett_crud':F_enterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def enterprises_people_dett(request):

    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('idpeo') is not None:
            enterprise_people_dett=EnterprisesPeople.objects.filter(id=request.GET.get('idpeo')).first()
            enterprise_dett=Enterprises.objects.filter(id=enterprise_people_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})

            TodosV=Todos.objects.filter(enterprisepeopleobj=enterprise_people_dett.id,type='todo')
            ReqV=Todos.objects.filter(enterprisepeopleobj=enterprise_people_dett.id,type='requirement')
            RolesV=Roles.objects.filter(enterprises_people=enterprise_people_dett.id)
            Roles_typeV=RolesType.objects.filter(tenants=current_tenant.id )
            orgchart_dett=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id )

            F_enterprises_people_dett_crud=Enterprises_people_crud(instance=enterprise_people_dett)
            TodosTemplateV=TodosTemplate.objects.all()
            return render(request, 'users/enterprises_people_dett.html', {'TodosTemplateV':TodosTemplateV,'Roles_typeV':Roles_typeV,'orgchart_dett':orgchart_dett,'RolesV':RolesV,'TodosV':TodosV,'ReqV':ReqV,'title':"Persone", 'F_enterprises_people_dett_crud':F_enterprises_people_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprises')).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 a=EnterprisesPeople(id=str(uuid.uuid4()),name=request.POST.get('name'),surname=request.POST.get('surname'),enterprises=enterprise_dett)
                 a.save()
                 enterprise_people_dett=EnterprisesPeople.objects.filter(id=a.id).first()
                 F_enterprises_people_dett_crud=Enterprises_people_crud(initial={'id': enterprise_people_dett.id,'name': enterprise_people_dett.name,'surname': enterprise_people_dett.surname,'enterprises': enterprise_people_dett.enterprises})
                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/enterprises_people_dett.html', {'title':"Persone", 'F_enterprises_people_dett_crud':F_enterprises_people_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            enterprise_people_dett=EnterprisesPeople.objects.filter(id=request.POST.get('id')).first()
            enterprise_dett=Enterprises.objects.filter(id=enterprise_people_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':EnterprisesV})
            form = Enterprises_people_crud(request.POST, instance=enterprise_people_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    enterprise_people_dett=EnterprisesPeople.objects.filter(id=request.POST.get('id')).first()
                    F_enterprises_people_dett_crud=Enterprises_people_crud(initial={'id': enterprise_people_dett.id,'surname': enterprise_people_dett.surname,'name': enterprise_people_dett.name,'enterprises': enterprise_people_dett.enterprises})
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/enterprises_people_dett.html', {'title':"Azienda",'F_enterprises_people_dett_crud':F_enterprises_people_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  enterprise_people_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Aggiungi azienda",'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
    F_enterprises_people_dett_crud=Enterprises_people_crud(initial={'enterprises': current_enterprise.id})

    return render(request, 'users/enterprises_people_dett.html', {'title':"Aggiungi persona",'F_enterprises_people_dett_crud':F_enterprises_people_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def orgcharts_dett(request):

    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('ido') is not None:
            orgchart_dett=OrgCharts.objects.filter(id=request.GET.get('ido')).first()
            enterprise_dett=Enterprises.objects.filter(id=orgchart_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})
            #F_orgcharts_dett_crud=Orgchart_crud(initial={'id': orgchart_dett.id,'name': orgchart_dett.name,'enterprises': orgchart_dett.enterprises})
            F_orgcharts_dett_crud=Orgchart_crud(instance=orgchart_dett)
            OrgChartsItems_dett=OrgChartsItems.objects.filter(orgcharts=orgchart_dett)
            return render(request, 'users/orgcharts_dett.html', {'OrgChartsItems_dett':OrgChartsItems_dett,'title':"Organigramma", 'F_orgcharts_dett_crud':F_orgcharts_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprises')).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 a=OrgCharts(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=enterprise_dett)
                 a.save()
                 orgchart_dett=OrgCharts.objects.filter(id=a.id).first()
                 F_orgcharts_dett_crud=Orgchart_crud(initial={'id': orgchart_dett.id,'name': orgchart_dett.name,'enterprises': orgchart_dett.enterprises})
                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/orgcharts_dett.html', {'title':"Organigramma", 'F_orgcharts_dett_crud':F_orgcharts_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            orgchart_dett=OrgCharts.objects.filter(id=request.POST.get('id')).first()
            enterprise_dett=Enterprises.objects.filter(id=orgchart_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':EnterprisesV})
            form = Orgchart_crud(request.POST, instance=orgchart_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    orgchart_dett=OrgCharts.objects.filter(id=request.POST.get('id')).first()
                    F_orgcharts_dett_crud=Orgchart_crud(initial={'id': orgchart_dett.id,'name': orgchart_dett.name,'enterprises': orgchart_dett.enterprises})
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/orgcharts_dett.html', {'title':"Organigramma",'F_orgcharts_dett_crud':F_orgcharts_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  orgchart_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
                  OrgChartsV=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Aggiungi organigramma",'OrgChartsV':OrgChartsV,'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    OrgChartsV=OrgCharts.objects.filter(enterprises__tenants=current_tenant.id)
    F_orgcharts_dett_crud=Orgchart_crud(initial={'enterprises': current_enterprise.id})
    return render(request, 'users/orgcharts_dett.html', {'title':"Aggiungi elemento organigramma",'F_orgcharts_dett_crud':F_orgcharts_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


@login_required()
def orgchartsitems_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()

    if request.method == "GET": #visualizzo
          if request.GET.get('action') =="New":
             OrgChartsV=OrgCharts.objects.filter(id=request.GET.get('id')).first()
             enterprise_dett=Enterprises.objects.filter(id=OrgChartsV.enterprises.id).first()
             if enterprise_dett not in EnterprisesV:
                   return render(request, 'users/vuota.html',{'a':EnterprisesV})



             F_OrgChartsItems_crud=OrgChartsItems_crud(initial={'orgcharts': OrgChartsV.id})
             return render(request, 'users/orgchartsitems_dett.html', {'title':"Nuova casella", 'F_OrgChartsItems_crud':F_OrgChartsItems_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

          if request.GET.get('action') =="Edit":
              OrgChartsItems_dett=OrgChartsItems.objects.filter(id=request.GET.get('id')).first()
              enterprise_dett=Enterprises.objects.filter(id=OrgChartsItems_dett.orgcharts.enterprises.id).first()
              if enterprise_dett not in EnterprisesV:
                  return render(request, 'users/vuota.html',{'a':EnterprisesV})

              F_OrgChartsItems_crud=OrgChartsItems_crud(instance=OrgChartsItems_dett)
              return render(request, 'users/orgchartsitems_dett.html', {'title':"Elemento organigramma", 'F_OrgChartsItems_crud':F_OrgChartsItems_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


    if request.method == "POST": #visualizzo


        if request.POST.get('action') =="Save":
           OrgChartsItems_dett=OrgChartsItems.objects.filter(id= request.POST.get('id')).first()
           OrgChartsV=OrgCharts.objects.filter(id=request.POST.get('orgcharts')).first()
           enterprise_dett=Enterprises.objects.filter(id=OrgChartsV.enterprises.id).first()
           if enterprise_dett not in EnterprisesV:
                   return render(request, 'users/vuota.html',{'a':EnterprisesV})


           form = OrgChartsItems_crud(request.POST, instance=OrgChartsItems_dett)

           if form.is_valid():
              instance=form.save()
              messages.success(request, "Salvato con successo")

           for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, error)
           OrgChartsItems_dett=OrgChartsItems.objects.filter(id= instance.id).first()
           F_OrgChartsItems_crud=OrgChartsItems_crud(instance=OrgChartsItems_dett)
           return render(request, 'users/orgchartsitems_dett.html', {'title':"Elemento organigramma", 'F_OrgChartsItems_crud':F_OrgChartsItems_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


        if request.POST.get('action')=="Delete":
                OrgChartsItems_dett=OrgChartsItems.objects.filter(id= request.POST.get('id')).first()
                enterprise_dett=Enterprises.objects.filter(id=OrgChartsItems_dett.orgcharts.enterprises.id).first()
                if enterprise_dett not in EnterprisesV:
                       return render(request, 'users/vuota.html',{'a':EnterprisesV})

                OrgChartsV=OrgCharts.objects.filter(id=OrgChartsItems_dett.orgcharts.id).first()
                OrgChartsItems_dett.delete()
                messages.success(request, "Cancellato con successo")

                return redirect('/orgcharts_dett/?ido='+ str(OrgChartsV.id))






@login_required()
def assets_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('ida') is not None:
            assets_dett=Assets.objects.filter(id=request.GET.get('ida')).first()
            enterprise_dett=Enterprises.objects.filter(id=assets_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})
            F_assets_dett_crud=Assets_crud(instance=assets_dett)
            TodosV=Todos.objects.filter(assets=assets_dett.id,type='todo')
            ReqV=Todos.objects.filter(assets=assets_dett.id,type='requirement')
            return render(request, 'users/assets_dett.html', {'TodosV':TodosV,'ReqV':ReqV,'title':"Asset", 'F_assets_dett_crud':F_assets_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprises')).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 a=Assets(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=enterprise_dett)
                 a.save()
                 assets_dett=Assets.objects.filter(id=a.id).first()
                 F_assets_dett_crud=Assets_crud(instance=assets_dett)
                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/assets_dett.html', {'title':"Asset", 'F_assets_dett_crud':F_assets_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            assets_dett=Assets.objects.filter(id=request.POST.get('id')).first()
            enterprise_dett=Enterprises.objects.filter(id=assets_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':EnterprisesV})
            form = Assets_crud(request.POST, instance=assets_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    assets_dett=Assets.objects.filter(id=request.POST.get('id')).first()
                    F_assets_dett_crud=Assets_crud(instance=assets_dett)
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/assets_dett.html', {'title':"Asset",'F_assets_dett_crud':F_assets_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  assets_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
                  AssetsV=Assets.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Asset",'AssetsV':AssetsV,'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    AssetsV=Assets.objects.filter(enterprises__tenants=current_tenant.id)
    F_assets_dett_crud=Assets_crud(instance=assets_dett)
    return render(request, 'users/assets_dett.html', {'AssetsV':AssetsV,'title':"Asset",'F_assets_dett_crud':F_assets_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def enterprises_enterprises_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('idee') is not None:
            enterprisesenterprises_dett=EnterprisesEnterprises.objects.filter(id=request.GET.get('idee')).first()
            enterprise_dett=Enterprises.objects.filter(id=enterprisesenterprises_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})

            F_enterprisesenterprises_dett_crud=EnterprisesEnterprises_crud(initial={'id': enterprisesenterprises_dett.id,'name': enterprisesenterprises_dett.name,'enterprises': enterprisesenterprises_dett.enterprises})
            return render(request, 'users/enterprises_enterprises_dett.html', {'title':"Persone", 'F_enterprisesenterprises_dett_crud':F_enterprisesenterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprises')).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 a=EnterprisesEnterprises(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=enterprise_dett)
                 a.save()
                 enterprisesenterprises_dett=EnterprisesEnterprises.objects.filter(id=a.id).first()
                 F_enterprisesenterprises_dett_crud=EnterprisesEnterprises_crud(initial={'id': enterprisesenterprises_dett.id,'name': enterprisesenterprises_dett.name,'enterprises': enterprisesenterprises_dett.enterprises})
                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/enterprises_enterprises_dett.html', {'title':"Persone giuridiche", 'F_enterprisesenterprises_dett_crud':F_enterprisesenterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            enterprisesenterprises_dett=EnterprisesEnterprises.objects.filter(id=request.POST.get('id')).first()
            enterprise_dett=Enterprises.objects.filter(id=enterprisesenterprises_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':EnterprisesV})
            form = EnterprisesEnterprises_crud(request.POST, instance=enterprisesenterprises_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    enterprisesenterprises_dett=EnterprisesEnterprises.objects.filter(id=request.POST.get('id')).first()
                    F_enterprisesenterprises_dett_crud=EnterprisesEnterprises_crud(initial={'id': enterprisesenterprises_dett.id,'name': enterprisesenterprises_dett.name,'enterprises': enterprisesenterprises_dett.enterprises})
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/enterprises_enterprises_dett.html', {'title':"Persone giuridiche",'F_enterprisesenterprises_dett_crud':F_enterprisesenterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  enterprisesenterprises_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
                  EnterprisesentErprisesAssetsV=EnterprisesEnterprises.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Persone giuridiche",'EnterprisesentErprisesAssetsV':EnterprisesentErprisesAssetsV,'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    EnterprisesentErprisesAssetsV=EnterprisesEnterprises.objects.filter(enterprises__tenants=current_tenant.id)
    F_enterprisesenterprises_dett_crud=EnterprisesEnterprises_crud(initial={'enterprises': current_enterprise.id})
    return render(request, 'users/enterprises_enterprises_dett.html', {'title':"Persone giuridiche",'F_enterprisesenterprises_dett_crud':F_enterprisesenterprises_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def enterprises_locations_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('idl') is not None:
            EnterprisesLocations_dett=EnterprisesLocations.objects.filter(id=request.GET.get('idl')).first()
            enterprise_dett=Enterprises.objects.filter(id=EnterprisesLocations_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})

            F_EnterprisesLocations_crud=EnterprisesLocations_crud(initial={'id': EnterprisesLocations_dett.id,'name': EnterprisesLocations_dett.name,'enterprises': EnterprisesLocations_dett.enterprises})
            return render(request, 'users/enterprises_locations_dett.html', {'title':"Persone", 'F_EnterprisesLocations_crud':F_EnterprisesLocations_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    if request.method == "POST":
            if request.POST.get('id')=="": #inserisco
                 enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprises')).first()
                 if enterprise_dett not in EnterprisesV:
                     return render(request, 'users/vuota.html',{'a':EnterprisesV})
                 a=EnterprisesLocations(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=enterprise_dett)
                 a.save()
                 EnterprisesLocations_dett=EnterprisesLocations.objects.filter(id=a.id).first()
                 F_EnterprisesLocations_crud=EnterprisesLocations_crud(initial={'id': EnterprisesLocations_dett.id,'name': EnterprisesLocations_dett.name,'enterprises': EnterprisesLocations_dett.enterprises})
                 messages.success(request, "Aggiunto con successo")
                 return render(request, 'users/enterprises_locations_dett.html', {'title':"Sedi", 'F_EnterprisesLocations_crud':F_EnterprisesLocations_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            EnterprisesLocations_dett=EnterprisesLocations.objects.filter(id=request.POST.get('id')).first()
            enterprise_dett=Enterprises.objects.filter(id=EnterprisesLocations_dett.enterprises.id).first()
            if enterprise_dett not in EnterprisesV:
                 return render(request, 'users/vuota.html',{'a':EnterprisesV})
            form = EnterprisesLocations_crud(request.POST, instance=EnterprisesLocations_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    EnterprisesLocations_dett=EnterprisesLocations.objects.filter(id=request.POST.get('id')).first()
                    F_EnterprisesLocations_crud=EnterprisesLocations_crud(initial={'id': EnterprisesLocations_dett.id,'name': EnterprisesLocations_dett.name,'enterprises': EnterprisesLocations_dett.enterprises})
                    messages.success(request, "Salvato con successo")
                    return render(request, 'users/enterprises_locations_dett.html', {'title':"Sedi",'F_EnterprisesLocations_crud':F_EnterprisesLocations_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

                if request.POST.get('azione')=="Delete":
                  #return render(request, 'users/vuota.html',{'a':request.POST})
                  EnterprisesLocations_dett.delete()
                  messages.success(request, "Cancellato con successo")
                  enterprises_peopleV=EnterprisesPeople.objects.filter(enterprises__tenants=current_tenant.id)
                  EnterprisesLocationsV=EnterprisesLocations.objects.filter(enterprises__tenants=current_tenant.id)
                  return render(request, 'users/base_tables.html',{'title':"Sedi",'EnterprisesLocationsV':EnterprisesLocationsV,'enterprises_peopleV':enterprises_peopleV,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

            for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, "Errore nel campo '{form[field].label}': {error}")
                return render(request, 'users/base_tables.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

    EnterprisesLocationsV=EnterprisesLocations.objects.filter(enterprises__tenants=current_tenant.id)
    F_EnterprisesLocations_crud=EnterprisesLocations_crud(initial={'enterprises': current_enterprise.id})
    return render(request, 'users/enterprises_locations_dett.html', {'title':"Sedi",'F_EnterprisesLocations_crud':F_EnterprisesLocations_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def actionsitemsstandards_dett(request):
     TypeOfActionsV=TypeOfActions.objects.all()
     StandardsV=Standards.objects.filter(state="Active")

     TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
     EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

     if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
           request.session['current_enterpriseid']=str(EnterprisesV.first().id)
           request.session['current_tenantid']=str(TenantsV.first().id)
     if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
           request.session['current_enterpriseid']=str(EnterprisesV.first().id)
           request.session['current_tenantid']=str(TenantsV.first().id)

     current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
     current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()



     if request.method == "POST":
                  Actions_dett=Actions.objects.filter(id=request.POST.get('current_action_id')).first()
                  Standard_item_dett=ItemsStandards.objects.filter(id=request.POST.get('standardid')).first()
                  if request.POST.get('azione')=="Add":


                      ActionsItemsStandards_dett=ActionsItemsStandards(id=str(uuid.uuid4()),items_standards=Standard_item_dett,name="",actions=Actions_dett)
                      ActionsItemsStandards_dett.save()
                      return redirect('/actions_dett/?idac='+request.POST.get('current_action_id'))
                  if request.POST.get('azione')=="Delete":
                    ActionsItemsStandards_dett=ActionsItemsStandards.objects.filter(id=request.POST.get('current_action_items_id')).first()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    ActionsItemsStandards_dett.delete()
                    messages.success(request, "Cancellato con successo")
                    #return render(request, 'users/actions_dett.html', {'title':"Azioni",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
                    return redirect('/actions_dett/?idac='+request.POST.get('current_action_id'))
     return render(request, 'users/vuota.html')

@login_required()
def actionsassets_dett(request):
     TypeOfActionsV=TypeOfActions.objects.all()
     StandardsV=Standards.objects.filter(state="Active")

     TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
     EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

     if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
           request.session['current_enterpriseid']=str(EnterprisesV.first().id)
           request.session['current_tenantid']=str(TenantsV.first().id)
     if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
           request.session['current_enterpriseid']=str(EnterprisesV.first().id)
           request.session['current_tenantid']=str(TenantsV.first().id)

     current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
     current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()



     if request.method == "POST":
                  Actions_dett=Actions.objects.filter(id=request.POST.get('current_action_id')).first()
                  Asset_dett=Assets.objects.filter(id=request.POST.get('current_assetid')).first()
                  if request.POST.get('azione')=="Add":
                      ActionsAssets_dett=ActionsAssets(id=str(uuid.uuid4()),assets=Asset_dett,name="",actions=Actions_dett)
                      ActionsAssets_dett.save()
                      return redirect('/actions_dett/?idac='+request.POST.get('current_action_id'))
                  if request.POST.get('azione')=="Delete":
                    ActionsAssets_dett=ActionsAssets.objects.filter(id=request.POST.get('current_asset_items_id')).first()
                    #return render(request, 'users/vuota.html',{'a':request.POST})
                    ActionsAssets_dett.delete()
                    messages.success(request, "Cancellato con successo")
                    #return render(request, 'users/actions_dett.html', {'title':"Azioni",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
                    return redirect('/actions_dett/?idac='+request.POST.get('current_action_id'))
     return render(request, 'users/vuota.html')

@login_required()
def enterprisestandards_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "GET": #visualizzo
        if request.GET.get('idstd') is not None:
            current_standards=Standards.objects.filter(id=request.GET.get('idstd')).first()
            enterprisestandards_dett=EnterpriseStandards.objects.filter(standards=request.GET.get('idstd')).first()
            if enterprisestandards_dett is None:
                 a=EnterpriseStandards(id=str(uuid.uuid4()),standards=current_standards,enterprises=current_enterprise)
                 a.save()
                 enterprisestandards_dett=EnterpriseStandards.objects.filter(id=a.id).first()

            #F_EnterpriseStandards_crud=EnterpriseStandards_crud(initial={'id': enterprisestandards_dett.id,'enterprises': enterprisestandards_dett.enterprises,'name': enterprisestandards_dett.name})
            F_EnterpriseStandards_crud=EnterpriseStandards_crud(instance=enterprisestandards_dett)
            TodosV=Todos.objects.filter(standards=current_standards,type='todo')
            ReqV=Todos.objects.filter(standards=current_standards,type='requirement')
            return render(request, 'users/enterprisestandards_dett.html', {'ReqV':ReqV,'TodosV':TodosV,'current_standards':current_standards,'title':current_standards.name, 'F_EnterpriseStandards_crud':F_EnterpriseStandards_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
    if request.method == "POST":
            enterprisestandards_dett=EnterpriseStandards.objects.filter(id=request.POST.get('id')).first()
            form = EnterpriseStandards_crud(request.POST, instance=enterprisestandards_dett)
            if form.is_valid():
                if request.POST.get('azione')=="Save":
                    form.save()
                    current_standards=Standards.objects.filter(id=request.POST.get('idstd')).first()
                    enterprisestandards_dett=EnterpriseStandards.objects.filter(id=request.POST.get('id')).first()
                    #F_EnterpriseStandards_crud=EnterpriseStandards_crud(initial={'id': enterprisestandards_dett.id,'name': enterprisestandards_dett.name,'enterprises': enterprisestandards_dett.enterprises})
                    F_EnterpriseStandards_crud=EnterpriseStandards_crud(instance=enterprisestandards_dett)
                    messages.success(request, "Salvato con successo")
                    TodosV=Todos.objects.filter(standards=current_standards,type='todo')
                    ReqV=Todos.objects.filter(standards=current_standards,type='requirement')
                    return render(request, 'users/enterprisestandards_dett.html', {'ReqV':ReqV,'TodosV':TodosV,'current_standards':enterprisestandards_dett,'title':enterprisestandards_dett.name, 'F_EnterpriseStandards_crud':F_EnterpriseStandards_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def enterprisesstandardsitems_dett(request):
   TypeOfActionsV=TypeOfActions.objects.all()
   StandardsV=Standards.objects.filter(state="Active")

   TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
   EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

   if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
         request.session['current_enterpriseid']=str(EnterprisesV.first().id)
         request.session['current_tenantid']=str(TenantsV.first().id)
   if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
         request.session['current_enterpriseid']=str(EnterprisesV.first().id)
         request.session['current_tenantid']=str(TenantsV.first().id)

   current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
   current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


   if request.method == "GET": #visualizzo
       if request.GET.get('idstdd') is not None:
           current_standards_item=ItemsStandards.objects.filter(id=request.GET.get('idstdd')).first()
           enterprisestandardsitems_dett=EnterprisesStandardsItems.objects.filter(standards_items=current_standards_item).first()
           if enterprisestandardsitems_dett is None:
                a=EnterprisesStandardsItems(id=str(uuid.uuid4()),standards_items=current_standards_item,enterprises=current_enterprise)
                a.save()
                enterprisestandardsitems_dett=EnterprisesStandardsItems.objects.filter(id=a.id).first()

           #F_EnterprisesStandardsItems_crud=EnterprisesStandardsItems_crud(initial={'id': enterprisestandardsitems_dett.id,'enterprises': enterprisestandardsitems_dett.enterprises,'name': enterprisestandardsitems_dett.name})
           F_EnterprisesStandardsItems_crud=EnterprisesStandardsItems_crud(instance=enterprisestandardsitems_dett)
           TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo',items_standards=current_standards_item)
           ReqV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='requirement',items_standards=current_standards_item)

           TOfAEnterprisesItemsStandards_dett=TOfAEnterprisesItemsStandards.objects.filter(enterprises_items_standards= enterprisestandardsitems_dett)

           return render(request, 'users/enterprisesstandardsitems_dett.html', {'TOfAEnterprisesItemsStandards_dett':TOfAEnterprisesItemsStandards_dett,'ReqV':ReqV,'TodosV':TodosV,'current_standards_item':current_standards_item,'title':current_standards_item.name, 'F_EnterprisesStandardsItems_crud':F_EnterprisesStandardsItems_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

   if request.method == "POST":
           enterprisestandardsitems_dett=EnterprisesStandardsItems.objects.filter(id=request.POST.get('id')).first()
           form = EnterpriseStandards_crud(request.POST, instance=enterprisestandardsitems_dett)
           if form.is_valid():
               if request.POST.get('azione')=="Save":
                   form.save()
                   enterprisestandardsitems_dett=EnterprisesStandardsItems.objects.filter(id=request.POST.get('id')).first()
                   F_EnterprisesStandardsItems_crud=EnterpriseStandards_crud(instance=enterprisestandardsitems_dett)
                   #F_EnterprisesStandardsItems_crud=EnterpriseStandards_crud(initial={'id': enterprisestandardsitems_dett.id,'name': enterprisestandardsitems_dett.name,'enterprises': enterprisestandardsitems_dett.enterprises})
                   messages.success(request, "Salvato con successo")
                   TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo',items_standards=enterprisestandardsitems_dett.standards_items)
                   ReqV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='requirement',items_standards=enterprisestandardsitems_dett.standards_items)
                   return render(request, 'users/enterprisesstandardsitems_dett.html', {'ReqV':ReqV,'TodosV':TodosV,'current_standards_item':enterprisestandardsitems_dett,'title':enterprisestandardsitems_dett.name , 'F_EnterprisesStandardsItems_crud':F_EnterprisesStandardsItems_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def todos_dett(request):
   TypeOfActionsV=TypeOfActions.objects.all()
   StandardsV=Standards.objects.filter(state="Active")

   TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
   EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

   if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
         request.session['current_enterpriseid']=str(EnterprisesV.first().id)
         request.session['current_tenantid']=str(TenantsV.first().id)
   if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
         request.session['current_enterpriseid']=str(EnterprisesV.first().id)
         request.session['current_tenantid']=str(TenantsV.first().id)

   current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
   current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()

   TodosV=Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo')
   if request.method == "GET": #visualizzo
       if request.GET.get('idt') is not None:
           todos_dett=Todos.objects.filter(id=request.GET.get('idt')).first()
           enterprise_dett=Enterprises.objects.filter(id=todos_dett.enterprises.id).first()
           if enterprise_dett not in EnterprisesV:
               return render(request, 'users/vuota.html',{'a':todos_dett})
           F_todos_dett_crud=Todos_crud(instance=todos_dett,  initial={'t':request.GET.get('t'),'ido':request.GET.get('ido'),'tod':request.GET.get('tod')} )
           return render(request, 'users/todos_dett.html', {'title':"Scadenzario", 'F_todos_dett_crud':F_todos_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

   if request.method == "POST":
           if request.POST.get('id')=="": #inserisco
                enterprise_dett=Enterprises.objects.filter(id=current_enterprise.id).first()
                #todo_template=request.POST.get('todo_template')
                if enterprise_dett not in EnterprisesV:
                    return render(request, 'users/vuota.html',{'a':EnterprisesV})
                todo_template_tmp=TodosTemplate.objects.filter(id=request.POST.get('todo_template')).first()

                if request.POST.get('t')=="Enterprises":
                    a=Todos(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=current_enterprise, todo_template=todo_template_tmp,type=request.POST.get('tod'),description=request.POST.get('description'), note=request.POST.get('note'),state=request.POST.get('state'), state_requirement=request.POST.get('state_requirement'))
                if request.POST.get('t')=="EnterprisesStandardItems":
                    tmp_std=ItemsStandards.objects.filter(id=request.POST.get('ido')).first()
                    a=Todos(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=current_enterprise, todo_template=todo_template_tmp,type=request.POST.get('tod'),state=request.POST.get('state'), state_requirement=request.POST.get('state_requirement'), items_standards=tmp_std)
                if request.POST.get('t')=="EnterprisesStandard":
                    tmp_std=Standards.objects.filter(id=request.POST.get('ido')).first()
                    a=Todos(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=current_enterprise,todo_template=todo_template_tmp, type=request.POST.get('tod'),state=request.POST.get('state'), state_requirement=request.POST.get('state_requirement'),  standards=tmp_std)
                if request.POST.get('t')=="EnterprisesPeople":
                    tmp_std=EnterprisesPeople.objects.filter(id=request.POST.get('ido')).first()
                    a=Todos(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=current_enterprise, todo_template=todo_template_tmp,type=request.POST.get('tod'),state=request.POST.get('state'), state_requirement=request.POST.get('state_requirement'),  enterprisepeopleobj=tmp_std)
                if request.POST.get('t')=="Assets":
                    tmp_std=Assets.objects.filter(id=request.POST.get('ido')).first()
                    a=Todos(id=str(uuid.uuid4()),name=request.POST.get('name'),enterprises=current_enterprise,todo_template=todo_template_tmp, type=request.POST.get('tod'),state=request.POST.get('state'), state_requirement=request.POST.get('state_requirement'),  assets=tmp_std)



                a.save()

                todos_dett=Todos.objects.filter(id=a.id).first()
                #F_todos_dett_crud=Todos_crud(instance=todos_dett)
                F_todos_dett_crud=Todos_crud(instance=todos_dett,  initial={'t':request.POST.get('t'),'ido':request.POST.get('ido'),'tod':request.POST.get('tod')} )
                messages.success(request, "Aggiunto con successo")
                return render(request, 'users/todos_dett.html', {'title':"Scadenzario", 'F_todos_dett_crud':F_todos_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

           todos_dett=Todos.objects.filter(id=request.POST.get('id')).first()
           enterprise_dett=Enterprises.objects.filter(id=todos_dett.enterprises.id).first()
           if enterprise_dett not in EnterprisesV:
                return render(request, 'users/vuota.html',{'a':EnterprisesV})
           form = Todos_crud(request.POST, instance=todos_dett)
           if form.is_valid():
               if request.POST.get('action')=="Save":


  #---------------------------------------------
                    from datetime import datetime, timedelta
                    from schedule.models import Event, Rule,Calendar


                    start_date = form.instance.recurrency_start
                    end_date = form.instance.recurrency_end

                    if form.instance.recurrency_enable=='YES':
                        calendar = Calendar.objects.first()
                        if not calendar:
                            new_calendar_name = "Calendario CFY"
                            slugg=str(uuid.uuid4())
                            calendar= Calendar(name=new_calendar_name, slug=slugg)
                            calendar.save()


                        if form.instance.recurrency_id_scheduler==0 or  form.instance.recurrency_id_scheduler is None :

                            rule = Rule(
                                name="Rule",
                                description="",
                                frequency=form.instance.recurrency_every,  # Imposta la frequenza su "MENSILE"
                                params=f"INTERVAL:{form.instance.recurrency_periodo}"
                            )
                            rule.save()

                            new_event = Event( title="->"+form.instance.name,  description=form.instance.description,start=start_date,   end=end_date,   rule=rule,calendar=calendar )
                            #new_event.end_recurring_period=form.instance.recurrency_end
                        else:

                            new_event = Event.objects.get(id=form.instance.recurrency_id_scheduler)
                            rule = new_event.rule

                            rule.frequency=form.instance.recurrency_every

                            rule.params= f"INTERVAL:{form.instance.recurrency_periodo}"

                            rule.save()


                            new_event.title = form.instance.name
                        #new_event.end_recurring_period=form.instance.recurrency_end
                            new_event.description = form.instance.description
                            new_event.start = start_date
                            new_event.end = end_date
                            new_event.rule = rule
                            new_event.calendar = calendar

                        new_event.save()

                        id_event=new_event.id
                        form.instance.recurrency_id_scheduler = id_event
                    else:
                        if  form.instance.recurrency_id_scheduler:
                            event_to_delete = Event.objects.get(id=form.instance.recurrency_id_scheduler)
                            ruleV = Rule.objects.get(id=event_to_delete.rule.id)
                            ruleV.delete()
                            event_to_delete.delete()
                            form.instance.recurrency_id_scheduler=0

#----------------------------------------------


                    form.save()
                    todos_dett=Todos.objects.filter(id=request.POST.get('id')).first()
                    F_todos_dett_crud=Todos_crud(instance=todos_dett)
                    messages.success(request, "Salvato con successo")


                    event_dett= Event.objects.filter(id=todos_dett.recurrency_id_scheduler).first()



                    start_date = datetime.now()
                    end_date = start_date + timedelta(days=365 * 10)  # Aggiungi 10 anni





                    if event_dett:
                        occ = event_dett.get_occurrences(start_date, end_date)
                    else:
                        occ = None







                    return render(request, 'users/todos_dett.html', {'occ':occ,'title':"Scadenzario",'F_todos_dett_crud':F_todos_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

               if request.POST.get('action')=="Delete":

                 todos_dett.delete()
                 messages.success(request, "Cancellato con successo")
                 return render(request, 'users/vuota.html')
           for field, error_list in form.errors.items():
               for error in error_list:
                   messages.error(request, error)
               return render(request, 'users/todos_dett.html',{'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
   F_todos_dett_crud=Todos_crud( initial={'todo_template':request.GET.get('temp'),'t':request.GET.get('t'),'ido':request.GET.get('ido'),'tod':request.GET.get('tod')} )



   title_todo=TodosTemplate.objects.filter(id=request.GET.get('temp')).first()
   return render(request, 'users/todos_dett.html', {'current_object_type':request.GET.get('t'), 'current_object_id':request.GET.get('ido'), 'type_of_todo':request.GET.get('tod'),'title':title_todo,'F_todos_dett_crud':F_todos_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

@login_required()
def actions_list(request):

         TypeOfActionsV=TypeOfActions.objects.all()
         StandardsV=Standards.objects.filter(state="Active")

         TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
         EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

         if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
               request.session['current_enterpriseid']=str(EnterprisesV.first().id)
               request.session['current_tenantid']=str(TenantsV.first().id)
         if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
               request.session['current_enterpriseid']=str(EnterprisesV.first().id)
               request.session['current_tenantid']=str(TenantsV.first().id)

         current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
         current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()

         current_actions_type=TypeOfActions.objects.filter(id=request.GET.get('idcat')).first()

         current_Actions = Actions.objects.filter(type_of_actions=current_actions_type, enterprises=current_enterprise)
         headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

         return render(request, 'users/actions.html', {'headbar':headbar,'MenuItem':"/",'title':"Aggiungi azienda",'current_actions_type':current_actions_type,'current_Actions':current_Actions, 'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})




@login_required()
def requirements_list(request):
           TypeOfActionsV=TypeOfActions.objects.all()
           StandardsV=Standards.objects.filter(state="Active")

           TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
           EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

           if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)
           if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)

           current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
           current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()
           headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

          # Ottieni i dati non filtrati
           todos = Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='requirement')

           # Istanzia il filtro e passa i dati non filtrati
           todos_filter = TodosFilter(request.GET, queryset=todos)

           # Applica il filtro ai dati
           todos_filtrati = todos_filter.qs

           # Il resto del tuo codice rimane invariato
           return render(request, 'users/requirements.html', {
               'headbar': headbar,
               'MenuItem': "Requirements",
               'title': "Requisiti",
               'ReqV': todos_filtrati,  # Passa i dati filtrati
               'current_tenant': current_tenant,
               'current_enterprise': current_enterprise,
               'Tenants': TenantsV,
               'TypeOfActions': TypeOfActionsV,
               'Enterprises': EnterprisesV,
               'Standards': StandardsV,
               'filterR': todos_filter,  # Passa il filtro al template
           })


@login_required()
def standarditems_list(request):
           TypeOfActionsV=TypeOfActions.objects.all()
           StandardsV=Standards.objects.filter(state="Active")

           TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
           EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

           if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)
           if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)

           current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
           current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()
           if request.GET.get('current_standardid') is not None:
               current_standard=Standards.objects.filter(id=request.GET.get('current_standardid')).first()
               current_ItemsStandards=ItemsStandards.objects.filter(standards=current_standard.id).order_by('order')

               headbar = [["Scadenze aperte", 12, "bg-warning"], ["Requisiti aperti", 12, "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

               StandardsItem_filter = StandardsItemFilter(request.GET, queryset=current_ItemsStandards)


               StandardsItem_filtrati = StandardsItem_filter.qs
               return render(request, 'users/standards.html', {'filterS': StandardsItem_filter,'headbar':headbar,'title':"Aggiungi azienda",'current_tenant':current_tenant,'current_standard':current_standard, 'current_ItemsStandards':StandardsItem_filtrati, 'current_enterprise':current_enterprise,'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})
           return render(request, 'users/standards.html', {'filterS': StandardsItem_filter,'headbar':headbar,'title':"Aggiungi azienda",'current_tenant':current_tenant,'current_standard':current_standard, 'current_ItemsStandards':StandardsItem_filtrati, 'current_enterprise':current_enterprise,'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

           return redirect('/')








# =============================================================================
# class TodosListView(FilterView):
#     model = Todos
#     filterset_class = TodosFilter
#     template_name = 'user/todos.html'
#     context_object_name = 'todos_lista'
# =============================================================================


@login_required()
def todos_list(request):
           TypeOfActionsV=TypeOfActions.objects.all()
           StandardsV=Standards.objects.filter(state="Active")

           TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
           EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

           if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)
           if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
                 request.session['current_enterpriseid']=str(EnterprisesV.first().id)
                 request.session['current_tenantid']=str(TenantsV.first().id)

           current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
           current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()
           headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]


           TodosTemplateV=TodosTemplate.objects.all()
           # Ottieni i dati non filtrati
           todos = Todos.objects.filter(enterprises=request.session['current_enterpriseid'],type='todo')

           # Istanzia il filtro e passa i dati non filtrati
           todos_filter = TodosFilter(request.GET, queryset=todos)

           # Applica il filtro ai dati
           todos_filtrati = todos_filter.qs

           # Il resto del tuo codice rimane invariato
           return render(request, 'users/todos.html', {
               'TodosTemplateV':TodosTemplateV,
               'headbar': headbar,
               'MenuItem': "Todos",
               'title': "Scadenzario",
               'TodosV': todos_filtrati,  # Passa i dati filtrati
               'current_tenant': current_tenant,
               'current_enterprise': current_enterprise,
               'Tenants': TenantsV,
               'TypeOfActions': TypeOfActionsV,
               'Enterprises': EnterprisesV,
               'Standards': StandardsV,
               'filterT': todos_filter,  # Passa il filtro al template
           })



# =============================================================================
# class TodosListView(FilterView):
#     model = Todos
#     filterset_class = TodosFilter
#     template_name = 'user/todos.html'
#     context_object_name = 'todos_lista'
# =============================================================================



@login_required()
def requests_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()

    if request.method == "GET": #visualizzo
          if request.GET.get('action') =="New":
              current_people=EnterprisesPeople.objects.filter(auth_user=request.user.id).first()
              F_requests_dett_crud=Requests_crud(initial={'enterprise': current_enterprise.id,'tenant': current_tenant.id,'whoami': current_people.id})
              return render(request, 'users/requests_dett.html', {'title':"Nuova richiesta", 'F_requests_dett_crud':F_requests_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})

          if request.GET.get('action') =="Edit":
              request_dett=Requests.objects.filter(id=request.GET.get('id')).first()
              enterprise_dett=Enterprises.objects.filter(id=request_dett.enterprise.id).first()
              if enterprise_dett not in EnterprisesV:
                  return render(request, 'users/vuota.html',{'a':EnterprisesV})

              F_requests_dett_crud=Requests_crud(instance=request_dett)
              return render(request, 'users/requests_dett.html', {'title':"Modifica richiesta", 'F_requests_dett_crud':F_requests_dett_crud,'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


    if request.method == "POST": #visualizzo
        if request.POST.get('action') =="Rmail":

            import  poplib
            Mailbox = poplib.POP3_SSL('pops.interhost.it', '995')
            Mailbox.user('comply@cply.eu')
            Mailbox.pass_('r_S_[k3.u97W2>@P2p')
            numMessages = len(Mailbox.list()[1])

            from email import policy
            from email.parser import BytesParser

            for i in range(numMessages):
                    raw_email = b"\n".join(Mailbox.retr(i+1)[1])
                    email_message = BytesParser(policy=policy.default).parsebytes(raw_email)

                    mittente = email_message.get("From")
                    destinatario = email_message.get("To")
                    oggetto = email_message.get("Subject")

                    if email_message.is_multipart():
                        for part in email_message.walk():
                            if part.get_content_type() == "text/plain":
                                corpo = part.get_payload(decode=True).decode(part.get_content_charset() or 'utf-8')
                    else:
                        corpo = email_message.get_payload(decode=True).decode(email_message.get_content_charset() or 'utf-8')


                    requestsV_dett=Requests(id=str(uuid.uuid4()),tenant=current_tenant,enterprise=current_enterprise,name=mittente+' '+oggetto, description=corpo, eml=raw_email)
                    requestsV_dett.save()








                    Mailbox.dele(i+1)

                    Mailbox.quit()
           # return render(request, 'users/vuota.html',{'a':numMessages})




        if request.POST.get('action') =="Save":
           request_dett=Requests.objects.filter(id= request.POST.get('id')).first()
           enterprise_dett=Enterprises.objects.filter(id=request.POST.get('enterprise')).first()
           if request.POST.get('tenant') != str(current_tenant.id):
                  return render(request, 'users/vuota.html',{'a':request.POST})
           if enterprise_dett not in EnterprisesV:
                  return render(request, 'users/vuota.html',{'a':enterprise_dett})

           form = Requests_crud(request.POST, instance=request_dett)

           if form.is_valid():
              form.save()
              messages.success(request, "Salvato con successo")

           for field, error_list in form.errors.items():
                for error in error_list:
                    messages.error(request, error)
           return redirect('/requests_dett/?action=Edit&id='+str(form.instance.id))
        if request.POST.get('action')=="Delete":
                request_dett=Requests.objects.filter(id= request.POST.get('id')).first()
                if request.POST.get('tenant') != str(current_tenant.id):
                       return render(request, 'users/vuota.html',{'a':EnterprisesV})
                request_dett.delete()
                messages.success(request, "Cancellato con successo")

    headbar = [["Attività non assegnate", "12", "bg-warning"], ["Attività aperte", "12", "bg-info"],["Requisiti non soddisfatti", "5", "bg-danger"],["Requisiti soddisfatti", "5", "bg-success"]]

    RequestsV=Requests.objects.filter(tenant=current_tenant.id)


    # Istanzia il filtro e passa i dati non filtrati
    requests_filter = RequestsFilter(request.GET, queryset=RequestsV)

    # Applica il filtro ai dati
    requests_filtrati = requests_filter.qs



    return render(request, 'users/requests.html',{'filterT': requests_filter,'headbar':headbar,'RequestsV':requests_filtrati,'MenuItem':"Requests",'title':"Richieste",'current_tenant':current_tenant,'current_enterprise':current_enterprise, 'Tenants':TenantsV,'TypeOfActions':TypeOfActionsV,'Enterprises':EnterprisesV,'Standards':StandardsV})


@login_required()
def roles_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "POST": #visualizzo
        if request.POST.get('action') =="Add":

                org_charts_itemsV=OrgChartsItems.objects.filter(id=request.POST.get('org_charts_items')).first()
                enterprises_peopleV=EnterprisesPeople.objects.filter(id=request.POST.get('enterprises_people')).first()
                roles_type_fkV=RolesType.objects.filter(id=request.POST.get('roles_type_fk')).first()

                roles_dett=Roles(id=str(uuid.uuid4()),name='',org_charts_items=org_charts_itemsV,enterprises_people=enterprises_peopleV,roles_type_fk=roles_type_fkV)
                roles_dett.save()
                return redirect('/enterprises_people_dett/?idpeo='+request.POST.get('enterprises_people'))


        if request.POST.get('action')=="Delete":
                roles_dett=Roles.objects.filter(id= request.POST.get('id')).first()
                if roles_dett.org_charts_items.orgcharts.enterprises not in EnterprisesV:
                       return render(request, 'users/vuota.html',{'a':EnterprisesV})
                roles_dett.delete()
                messages.success(request, "Cancellato con successo")
                return redirect('/enterprises_people_dett/?idpeo='+request.POST.get('enterprises_people'))

    return redirect('home')

@login_required()
def tofaenterprisesitemsstandards_dett(request):
    TypeOfActionsV=TypeOfActions.objects.all()
    StandardsV=Standards.objects.filter(state="Active")

    TenantsV=Tenants.objects.filter(authusertenants__auth_user=request.user.id)
    EnterprisesV=Enterprises.objects.filter(tenants=TenantsV.first().id)

    if 'current_tenantid' not in request.session or 'current_enterpriseid' not in request.session :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)
    if request.session['current_tenantid']=='' or request.session['current_enterpriseid']==''  :
          request.session['current_enterpriseid']=str(EnterprisesV.first().id)
          request.session['current_tenantid']=str(TenantsV.first().id)

    current_tenant=Tenants.objects.filter(id=request.session['current_tenantid']).first()
    current_enterprise=Enterprises.objects.filter(id=request.session['current_enterpriseid']).first()


    if request.method == "POST": #visualizzo
        #return render(request, 'users/vuota.html',{'a':request.POST})
        if request.POST.get('action') =="Add":

                type_of_actionsV=TypeOfActions.objects.filter(id=request.POST.get('type_of_actions')).first()
                enterprises_items_standardsV=EnterprisesStandardsItems.objects.filter(id=request.POST.get('enterprises_items_standards')).first()

                TOfAEnterprisesItemsStandards_dett=TOfAEnterprisesItemsStandards(id=str(uuid.uuid4()),name=request.POST.get('name'),tenant=current_tenant,enterprises_items_standards=enterprises_items_standardsV,type_of_actions=type_of_actionsV)
                TOfAEnterprisesItemsStandards_dett.save()

                return redirect('/enterprisesstandardsitems_dett/?idstdd='+request.POST.get('current_standards_item'))


        if request.POST.get('action')=="Delete":
                TOfAEnterprisesItemsStandards_dett=TOfAEnterprisesItemsStandards.objects.filter(id= request.POST.get('id')).first()
                if TOfAEnterprisesItemsStandards_dett.tenant not in TenantsV:
                       return render(request, 'users/vuota.html',{'a':EnterprisesV})
                TOfAEnterprisesItemsStandards_dett.delete()
                messages.success(request, "Cancellato con successo")

                return redirect('/enterprisesstandardsitems_dett/?idstdd='+request.POST.get('current_standards_item'))

    return redirect('home')
